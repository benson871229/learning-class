# -*- coding: utf-8 -*-
"""
產生 Word 報價／繳費通知單「範本」(docxtpl 樣板)。

執行後會在 範本/報價單範本.docx 產生一份可用 Word 直接開啟編輯的樣板檔，
裡面含 Jinja 佔位符 ({{ }}) 與可重複列 ({%tr%})，交給 產生報價單.py 套版。

版面完全對照使用者提供的格式：
    親愛的家長您好：
    {學生}小朋友報名課程如下：
    [課程名稱 | 日期起訖 | 學費 | 教材費 | 扣除金額 | 總計] 表格
    小計列 / 備註列
    繳費方式如下： 1. 現場繳費  2. 線上匯款 (帳戶資訊)

想改版面？直接用 Word 打開 範本/報價單範本.docx 修改即可，不必動這支程式；
若想讓「重建範本」可重現，改完設計邏輯後再跑這支程式覆蓋即可。
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

CN_FONT = "標楷體"  # 傳統楷書，貼近原稿；系統缺字時 Word 會自動替代


def set_cn_font(run, name=CN_FONT, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_para(doc, text, *, bold=False, size=12, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_cn_font(p.add_run(text), size=size, bold=bold)
    return p


def cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    set_cn_font(p.add_run(text), size=size, bold=bold)


def build():
    doc = Document()

    # 頁面預設字型
    style = doc.styles["Normal"]
    style.font.name = CN_FONT
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    # ── 開頭問候 ──
    add_para(doc, "親愛的家長您好：", size=12, space_after=6)

    # {學生}小朋友報名課程如下：  ── 學生姓名粗體
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_cn_font(p.add_run("{{ student_name }}"), size=12, bold=True)
    set_cn_font(p.add_run("小朋友報名課程如下："), size=12, bold=False)

    # ── 課程表格 ──
    headers = ["課程名稱", "日期起訖", "學費", "教材費", "扣除金額", "總計"]
    # 列：表頭 + [for 輔助列 + 可重複資料列 + endfor 輔助列] + 小計 + 備註
    # docxtpl {%tr%} 迴圈：for 與 endfor 各自獨立一列，中間資料列會依 courses 重複，
    # 兩條輔助列在套版後會被自動移除。
    table = doc.add_table(rows=6, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 欄寬（cm），總寬 ~16.5cm 貼近 A4 版心
    widths = [3.0, 4.0, 2.0, 2.0, 2.5, 2.5]

    # 表頭
    for j, h in enumerate(headers):
        cell_text(table.cell(0, j), h, bold=True)

    # for 輔助列
    cell_text(table.rows[1].cells[0], "{%tr for c in courses %}",
              align=WD_ALIGN_PARAGRAPH.LEFT)

    # 可重複資料列
    row = table.rows[2]
    cell_text(row.cells[0], "{{ c.name }}")
    cell_text(row.cells[1], "{{ c.date }}")
    cell_text(row.cells[2], "{{ c.tuition }}")
    cell_text(row.cells[3], "{{ c.material }}")
    cell_text(row.cells[4], "{{ c.deduction }}")
    cell_text(row.cells[5], "{{ c.total }}")

    # endfor 輔助列
    cell_text(table.rows[3].cells[0], "{%tr endfor %}",
              align=WD_ALIGN_PARAGRAPH.LEFT)

    # 小計列：前兩格合併顯示「小 計」
    sub = table.rows[4]
    a = sub.cells[0].merge(sub.cells[1])
    cell_text(a, "小　計", bold=True)
    cell_text(sub.cells[2], "{{ sum_tuition }}", bold=True)
    cell_text(sub.cells[3], "{{ sum_material }}", bold=True)
    cell_text(sub.cells[4], "{{ sum_deduction }}", bold=True)
    cell_text(sub.cells[5], "{{ sum_total }}", bold=True)

    # 備註列：整列合併
    note = table.rows[5]
    merged = note.cells[0]
    for k in range(1, 6):
        merged = merged.merge(note.cells[k])
    cell_text(merged, "備註：{{ note }}", align=WD_ALIGN_PARAGRAPH.LEFT)

    # 設定欄寬
    for r in table.rows:
        for j, w in enumerate(widths):
            if j < len(r.cells):
                r.cells[j].width = Cm(w)

    # ── 繳費方式（固定文字，可直接在 Word 編輯）──
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_para(doc, "繳費方式如下：", bold=True, space_after=4)

    add_para(doc, "1.　現場繳費", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "2.　線上匯款(轉帳完成請告知帳戶末五碼)",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "　　　匯款帳戶：　FUN 學院文理補習班",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "　　　銀行代號：　812(台新銀行)",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "　　　匯款帳號：　21060100208180　　　　繳費完成記得領取收據！",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    out = Path(__file__).parent / "範本" / "報價單範本.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"範本已產生：{out}")


if __name__ == "__main__":
    build()
