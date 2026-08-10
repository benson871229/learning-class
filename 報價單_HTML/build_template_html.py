# -*- coding: utf-8 -*-
"""
產生「HTML 版」專用的 Word 範本（B5，docxtemplater 語法）。

與 Python 版（報價系統/build_template.py）的差別只在標記語法：
    Python 版 docxtpl        →  {{ student_name }} 、 {%tr for c in courses %}
    HTML 版 docxtemplater    →  {student_name}    、 {#courses} … {/courses}

版面（B5 橫式 257×182mm）與內容完全一致。

執行：python3 build_template_html.py
輸出：範本/報價單範本.docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

CN_FONT = "標楷體"


def set_cn_font(run, name=CN_FONT, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_para(doc, text, *, bold=False, size=12, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_cn_font(p.add_run(text), size=size, bold=bold)
    return p


def cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    set_cn_font(p.add_run(text), size=size, bold=bold)


def build():
    doc = Document()

    # ── 頁面：B5（JIS）橫式 257 × 182 mm ──
    # python-docx 的 orientation 只是標記，實際尺寸仍要自己指定（寬 > 高）
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width = Mm(257)
    s.page_height = Mm(182)
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = CN_FONT
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    # ── 開頭 ──
    add_para(doc, "親愛的家長您好：")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_cn_font(p.add_run("{student_name}"), size=12, bold=True)
    set_cn_font(p.add_run("小朋友報名課程如下："), size=12)

    # ── 課程表格 ──
    headers = ["課程名稱", "日期起訖", "學費", "教材費", "扣除金額", "總計"]
    table = doc.add_table(rows=4, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 欄寬（cm）合計 22.1，正好填滿橫式 B5 版心（頁寬 25.7 − 左右邊界各 1.8）
    widths = [4.5, 5.5, 2.8, 2.8, 3.2, 3.3]

    for j, h in enumerate(headers):
        cell_text(table.cell(0, j), h, bold=True)

    # 可重複列：docxtemplater 的 {#courses} … {/courses} 放在同一列的頭尾格
    row = table.rows[1]
    cell_text(row.cells[0], "{#courses}{name}")
    cell_text(row.cells[1], "{date}")
    cell_text(row.cells[2], "{tuition}")
    cell_text(row.cells[3], "{material}")
    cell_text(row.cells[4], "{deduction}")
    cell_text(row.cells[5], "{total}{/courses}")

    # 小計
    sub = table.rows[2]
    a = sub.cells[0].merge(sub.cells[1])
    cell_text(a, "小　計", bold=True)
    cell_text(sub.cells[2], "{sum_tuition}", bold=True)
    cell_text(sub.cells[3], "{sum_material}", bold=True)
    cell_text(sub.cells[4], "{sum_deduction}", bold=True)
    cell_text(sub.cells[5], "{sum_total}", bold=True)

    # 備註（整列合併）
    note = table.rows[3]
    merged = note.cells[0]
    for k in range(1, 6):
        merged = merged.merge(note.cells[k])
    cell_text(merged, "備註：{note}", align=WD_ALIGN_PARAGRAPH.LEFT)

    for r in table.rows:
        for j, w in enumerate(widths):
            if j < len(r.cells):
                r.cells[j].width = Cm(w)

    # 關閉自動調整，並讓 tblGrid 與實際欄寬一致。
    # 只設 cell.width 的話 tblGrid 仍是平均分配，Word 可能照平均值畫成等寬六欄。
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
    tblPr.append(layout)
    for gridcol, w in zip(table._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol")), widths):
        gridcol.set(qn("w:w"), str(int(round(Cm(w).twips))))

    # ── 繳費方式 ──
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_para(doc, "繳費方式如下：", bold=True, space_after=4)
    add_para(doc, "1.　現場繳費", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "2.　線上匯款(轉帳完成請告知帳戶末五碼)",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "　　　匯款帳戶：　FUN 學院文理補習班",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "　　　銀行代號：　812(台新銀行)",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "　　　匯款帳號：　21060100208180　　　　繳費完成記得領取收據！",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    out = Path(__file__).parent / "範本" / "報價單範本.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"HTML 版範本已產生（B5）：{out}")


if __name__ == "__main__":
    build()
